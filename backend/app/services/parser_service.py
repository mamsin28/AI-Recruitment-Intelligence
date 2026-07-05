# backend/app/services/parser_service.py

from io import BytesIO
from pathlib import Path

from app.exceptions.parser_exceptions import (
    ParserException,
    ParserServiceException,
)
# from app.services.parser_service import ParserService

import fitz
from docx import Document

from app.config.logging_config import get_logger
from app.models.parsed_document import ParsedDocument

logger = get_logger(__name__)


# class ParserException(Exception):
#     """Raised when document parsing fails."""

class PDFParser:
    """Parser for PDF documents."""

    def parse(self, content: bytes) -> tuple[str, int]:
        text = []

        try:
            with fitz.open(stream=content, filetype="pdf") as pdf:
                logger.info("PDF opened successfully.")

                page_count = len(pdf)

                for page in pdf:
                    text.append(page.get_text())

                logger.info(
                    "PDF parsed successfully (%s pages).",
                    page_count,
                )

            return "\n".join(text).strip(), page_count

        # except Exception as exc:
        #     logger.exception("Failed to parse PDF document.")
        #     raise ParserException("Unable to parse PDF document.") from exc

        except ParserException:
            logger.exception("ParserService failed, Failed to parse PDF document..")
            raise

class DOCXParser:
    """Parser for Microsoft Word documents."""

    def parse(self, content: bytes) -> tuple[str, int]:
        try:
            document = Document(BytesIO(content))

            text = "\n".join(
                paragraph.text
                for paragraph in document.paragraphs
            ).strip()

            logger.info("DOCX parsed successfully.")

            # Word documents لا تحتوي على مفهوم الصفحات أثناء القراءة
            return text, 1

        except Exception as exc:
            logger.exception("Failed to parse DOCX document.")
            raise ParserException("Unable to parse DOCX document.") from exc


class ParserService:
    """
    Service responsible for selecting the correct parser
    based on the uploaded document type.
    """

    def __init__(self) -> None:
        self._parsers = {
            ".pdf": PDFParser(),
            ".docx": DOCXParser(),
        }

    def parse(
        self,
        filename: str,
        content: bytes,
    ) -> ParsedDocument:

        extension = Path(filename).suffix.lower()

        logger.info("Selecting parser for '%s'.", extension)

        parser = self._parsers.get(extension)

        if parser is None:
            logger.error("Unsupported file type: %s", extension)
            raise ParserServiceException(
                f"Unsupported file type: {extension}"
            )

        logger.info("Parsing started for '%s'.", filename)

        text, page_count = parser.parse(content)

        document = ParsedDocument(
            filename=filename,
            extension=extension,
            text=text,
            page_count=page_count,
            character_count=len(text),
            word_count=len(text.split()),
        )

        logger.info(
            "Parsing completed successfully. %s words extracted.",
            document.word_count,
        )

        return document